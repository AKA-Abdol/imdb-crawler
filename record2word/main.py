from docx import Document
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH as wap
from docx.shared import Inches, RGBColor

INDEX = 'ردیف'
CONTEXT = 'مقوله'
DATA = 'اطلاعات'

table_translation = {
    'genres': 'ژانر',
    'rating': 'امتیاز',
    'duration':'',
    'seasons':'تعداد فصل',
    'episodes':'تعداد قسمت',
    'stars':'پرسوناژ اصلی'
}

csv_dir = input('directory of csv to read: ')
count = int(input('how many to generate: '))
df = pd.read_csv(csv_dir)

for i in range(0, count):
    doc = Document()
    record = df.loc[i]
    title = record['title'].replace('\n', '').replace(':', '')
    # intro
    doc.add_heading(f"{title.split('.')[1].strip()} / ({record['year']})".replace('\n', ' '), level=1)
    doc.add_paragraph(record['content'].replace('\n', ''))
    link_paragraph = doc.add_paragraph()
    link_paragraph.add_run(record['link']).font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    # table
    table = doc.add_table(1, 3)
    table.style = 'Table Grid'
    table.alignment = wap.CENTER
    heading_cells = table.rows[0].cells
    heading_cells[2].text = INDEX
    heading_cells[1].text = CONTEXT
    heading_cells[0].text = DATA
    for idx, context in enumerate(['genres', 'rating', 'duration', 'seasons', 'episodes', 'stars']):
        cells = table.add_row().cells 
        cells[2].text = f'{idx + 1}.'
        cells[1].text = table_translation[context]
        data = f'{record[context]}'.replace('\n', ' ')
        cells[0].text = data.split('.')[0] if context in ['seasons', 'episodes'] else data
    cells = table.add_row().cells
    cells[2].text = '7.'
    cells[1].text = 'کمپانی'
    cells = table.add_row().cells
    cells[2].text = '8.'
    cells[1].text = 'تعداد قسمت دیده شده'
    for cell in table.columns[2].cells:
        cell.width = Inches(0.5)
    doc.save(f".words/{title}.docx")
