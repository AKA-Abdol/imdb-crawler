from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd


table_translation = {
    "genres": "ژانر",
    "rating": "امتیاز",
    "duration": "",
    "seasons": "تعداد فصل",
    "episodes": "تعداد قسمت",
    "stars": "پرسوناژ اصلی",
}


def add_existing_doc_at_end(doc, existing_doc):
    for element in existing_doc.element.body:
        doc.element.body.append(element)

    return doc


# Function to add styled paragraphs at the beginning of a document
def add_paragraphs(doc, paragraphs, direction="rtl"):
    # Add new styled paragraphs
    for paragraph_text, style in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(paragraph_text)

        # Apply styles
        if "font_name" in style:
            run.font.name = style["font_name"]
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), style["font_name"])

        if "font_size" in style:
            run.font.size = Pt(style["font_size"])

        if "font_color" in style:
            run.font.color.rgb = RGBColor(*style["font_color"])

        if "bold" in style:
            run.bold = style["bold"]

        if "italic" in style:
            run.italic = style["italic"]

        if "centered" in style:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Set text direction
        if direction == "rtl":
            pPr = p._element.get_or_add_pPr()
            bidi = OxmlElement("w:bidi")
            pPr.append(bidi)
        elif direction == "ltr":
            pPr = p._element.get_or_add_pPr()
            bidi = pPr.find(qn("w:bidi"))
            if bidi is not None:
                pPr.remove(bidi)

    return doc


def add_styled_text(p, run, style, direction = 'rtl'):
    if "font_name" in style:
        run.font.name = style["font_name"]
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), style["font_name"])

    if "font_size" in style:
        run.font.size = Pt(style["font_size"])

    if "font_color" in style:
        run.font.color.rgb = RGBColor(*style["font_color"])

    if "bold" in style:
        run.bold = style["bold"]

    if "italic" in style:
        run.italic = style["italic"]
    
    # Set text direction
    if direction == "rtl":
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    elif direction == "ltr":
        pPr = p._element.get_or_add_pPr()
        bidi = pPr.find(qn("w:bidi"))
        if bidi is not None:
            pPr.remove(bidi)


# Function to add a styled table to a document
def add_styled_table(doc, table_data):
    table = doc.add_table(rows=1, cols=len(table_data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    # Adding the header row with styles
    for i, (header, style) in enumerate(table_data[0]):
        hdr_paragraph = hdr_cells[i].paragraphs[0]
        hdr_run = hdr_paragraph.add_run(header)
        add_styled_text(hdr_paragraph, hdr_run, style)

    # Adding the rest of the table data with styles
    for row_data in table_data[1:]:
        row_cells = table.add_row().cells
        for i, (cell_text, style) in enumerate(row_data):
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_run = cell_paragraph.add_run(cell_text)
            add_styled_text(cell_paragraph, cell_run, style)

    return doc


csv_path = "./delivery.csv"
df = pd.read_csv(csv_path)

for i in range(2):
    # Load an existing Word document or create a new one
    doc = Document()
    record = df.loc[i]
    title = record["title"].replace("\n", "").replace(":", "")
    new_paragraphs = [
        (
            "فرم تحلیل برنامه ها",
            {
                "font_name": "Baghdad",
                "font_size": 20,
                "font_color": (0, 0, 0),  # RGB (Red, Green, Blue)
                "bold": True,
                "italic": False,
                "centered": True
            }
        ),
        (
            f"{title.split('.')[1].strip()} / ({record['year']})".replace("\n", " "),
            {
                "font_name": "Baghdad",
                "font_size": 16,
                "font_color": (0, 0, 0),  # RGB (Red, Green, Blue)
                "bold": True,
                "italic": False,
            },
        ),
        (
            record["fa_logline"].replace("\n", ""),
            {
                "font_name": "Baghdad",
                "font_size": 16,
                "font_color": (0, 0, 0),  # RGB (Red, Green, Blue)
                "bold": False,
                "italic": False,
            },
        ),
        (
            record["link"],
            {
                "font_name": "Arial",
                "font_size": 16,
                "font_color": (0, 0, 255),  # RGB (Red, Green, Blue)
                "bold": True,
                "italic": True,
            },
        ),
    ]

    doc = add_paragraphs(doc, new_paragraphs)

    # table
    table_data = [
        [
            (
                "اطلاعات",
                {
                    "font_name": "Baghdad",
                    "font_size": 16,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
            (
                "مقوله",
                {
                    "font_name": "Baghdad",
                    "font_size": 16,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
            (
                "ردیف",
                {
                    "font_name": "Baghdad",
                    "font_size": 16,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
        ],
    ]

    for idx, context in enumerate(
        ["genres", "rating", "duration", "seasons", "episodes", "stars"]
    ):
        data = f"{record[context]}".replace("\n", " ")
        table_data.append(
            [
                (
                    data.split(".")[0] if context in ["seasons", "episodes"] else data,
                    {
                        "font_name": "Times New Roman",
                        "font_size": 16,
                        "font_color": (0, 0, 0),
                        "bold": False,
                        "italic": False,
                    },
                ),
                (
                    table_translation[context],
                    {
                        "font_name": "Baghdad",
                        "font_size": 14,
                        "font_color": (0, 0, 0),
                        "bold": False,
                        "italic": False,
                    },
                ),
                (
                    f"{idx + 1}.",
                    {
                        "font_name": "Baghdad",
                        "font_size": 14,
                        "font_color": (0, 0, 0),
                        "bold": False,
                        "italic": False,
                    },
                ),
            ]
        )

    table_data.append(
        [
            (
                "",
                {
                    "font_name": "Times New Roman",
                    "font_size": 16,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
            (
                "کمپانی",
                {
                    "font_name": "Baghdad",
                    "font_size": 14,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
            (
                "7.",
                {
                    "font_name": "Baghdad",
                    "font_size": 14,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
        ]
    )

    table_data.append(
        [
            (
                "",
                {
                    "font_name": "Times New Roman",
                    "font_size": 16,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
            (
                "تعداد قسمت دیده شده",
                {
                    "font_name": "Baghdad",
                    "font_size": 14,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
            (
                "8.",
                {
                    "font_name": "Baghdad",
                    "font_size": 14,
                    "font_color": (0, 0, 0),
                    "bold": False,
                    "italic": False,
                },
            ),
        ]
    )
    add_styled_table(doc, table_data)
    existing_doc_path = "./record2word/Form.docx"
    existing_doc = Document(existing_doc_path)
    add_existing_doc_at_end(doc, existing_doc)
    doc.save(f".words/{title}.docx")
